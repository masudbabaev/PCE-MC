import pandas as pd
import numpy as np
import chaospy as cp
import matplotlib.pyplot as plt
import seaborn as sns
import os
from docx import Document  # For Word document export
import traceback  # For better error reporting

# Print ChaosPy version for debugging
print(f"ChaosPy version: {cp.__version__}")

class PCEvsMC:
    def __init__(self, file_path: str, input_file: str, output_dir: str = "output_plots"):
        self.file_path = file_path
        self.input_file = input_file
        self.output_folder = os.path.join(file_path, output_dir)
        os.makedirs(self.output_folder, exist_ok=True)
        self.input_params = ['BSORW', 'AQVISC', 'PERMI', 'POR', 'PERMK']
        self.outputs = ['OilRecoveryFactor', 'WaterCut', 'OIIP']
        self.max_pce_degree = 6

    def load_data(self) -> pd.DataFrame:
        """Load the dataset."""
        full_path = os.path.join(self.file_path, self.input_file)
        df = pd.read_csv(full_path)
        df = df[self.input_params + self.outputs].dropna()
        # Remove rows with infinite values
        df = df.replace([np.inf, -np.inf], np.nan).dropna()
        print(f"Data loaded. Shape: {df.shape}")
        return df

    def get_distributions(self) -> cp.J:
        """Define input distributions, treating all variables as independent (temporary)."""
        means = {
            'BSORW': 0.3,
            'AQVISC': 0.5,
            'PERMI': 1.0,
            'POR': 1.0,
            'PERMK': 1.0
        }
        stds = {
            'BSORW': 0.0375,
            'AQVISC': 0.0125,
            'PERMI': 0.125,
            'POR': 0.125,
            'PERMK': 0.125
        }

        # Define independent normal distributions for all parameters
        bsorw_dist = cp.Normal(means['BSORW'], stds['BSORW'])
        aqvisc_dist = cp.Normal(means['AQVISC'], stds['AQVISC'])
        permi_dist = cp.Normal(means['PERMI'], stds['PERMI'])
        por_dist = cp.Normal(means['POR'], stds['POR'])
        permk_dist = cp.Normal(means['PERMK'], stds['PERMK'])

        # Combine all distributions into a joint distribution
        joint_dist = cp.J(bsorw_dist, aqvisc_dist, permi_dist, por_dist, permk_dist)
        print(f"Joint distribution variables: {[str(dist) for dist in joint_dist]}")
        return joint_dist

    def build_pce_model(self, X: np.ndarray, Y: np.ndarray, joint_dist: cp.Distribution, degree: int) -> list:
        """Build PCE models for the given degree."""
        poly = cp.generate_expansion(degree, joint_dist)
        pce_models = []
        for i, out in enumerate(self.outputs):
            model = cp.fit_regression(poly, X.T, Y[i])
            print(f"Built PCE model for {out} at degree {degree}")
            pce_models.append(model)
        return pce_models, poly

    def calculate_relative_error(self, y_true: np.ndarray, y_pred: np.ndarray) -> float:
        """Calculate relative error in percentage: (MAE / mean(MC)) * 100."""
        valid = np.isfinite(y_true) & np.isfinite(y_pred)
        if valid.sum() < 2:
            return np.nan
        mae = np.mean(np.abs(y_true[valid] - y_pred[valid]))
        mean_true = np.mean(y_true[valid])
        if abs(mean_true) < 1e-6:
            return np.nan
        return (mae / mean_true) * 100

    def extract_coefficients_and_polynomial(self, pce_models, poly, degree, specific_output=None):
        """Extract coefficients, orders, and export the full polynomial to Excel and Word."""
        multi_indices = poly.exponents  # Use poly.exponents instead of poly.keys
        assert multi_indices.shape[1] == len(self.input_params), f"Expected {len(self.input_params)} variables, got {multi_indices.shape[1]}"
        print(f"multi_indices type: {type(multi_indices)}, shape: {multi_indices.shape}")
        print(f"First multi_index: {multi_indices[0]}, type of first element: {type(multi_indices[0][0])}")

        all_polynomial_data = []
        outputs_to_process = [specific_output] if specific_output else self.outputs
        models_to_process = pce_models if specific_output else pce_models

        for i, (model, out) in enumerate(zip(models_to_process, outputs_to_process)):
            coeffs = model.coefficients
            assert len(coeffs) == len(multi_indices), f"Mismatch for {out}: {len(coeffs)} coeffs vs {len(multi_indices)} terms"

            polynomial_terms = []
            polynomial_data = []

            for j, (coeff, multi_index) in enumerate(zip(coeffs, multi_indices)):
                if abs(coeff) < 1e-10:  # Skip negligible coefficients
                    continue

                # multi_index is a numpy array of integers, e.g., array([0, 0, 0, 0, 0])
                term = f"{coeff:.4e}"
                basis = []
                for k, (deg, param) in enumerate(zip(multi_index, self.input_params)):
                    if deg > 0:
                        mean, std = (0.3, 0.0375) if param == 'BSORW' else (0.5, 0.0125) if param == 'AQVISC' else (1.0, 0.125)
                        standardized = f"(({param} - {mean})/{std})"
                        basis.append(f"H_{deg}{standardized}")
                if not basis:
                    term += " * H_0"
                else:
                    term += " * " + " * ".join(basis)

                polynomial_terms.append(term)
                term_data = {
                    'output': out,
                    'term_index': j,
                    'coefficient': coeff,
                    'multi_index': str(tuple(multi_index)),  # Convert numpy array to tuple, then string
                    'term': term
                }
                for k, param in enumerate(self.input_params):
                    term_data[param + '_order'] = multi_index[k]  # Integer directly from numpy array
                polynomial_data.append(term_data)

            all_polynomial_data.extend(polynomial_data)

            # Print to console
            print(f"\n=== PCE Polynomial for {out} (Degree {degree}) ===")
            print("Term Index | Coefficient    | Multi-Index (BSORW, AQVISC, PERMI, POR, PERMK) | Polynomial Term")
            print("-" * 90)
            for data in polynomial_data:
                print(f"{data['term_index']:<10} | {data['coefficient']:<14.4e} | {data['multi_index']:<35} | {data['term']}")
            print(f"\nFull Polynomial for {out}:")
            full_poly = " + ".join(polynomial_terms) if polynomial_terms else "0"
            print(full_poly)

        # Export to CSV
        try:
            columns = ['output', 'term_index', 'coefficient', 'multi_index'] + \
                      [param + '_order' for param in self.input_params] + ['term']
            coeff_df = pd.DataFrame(all_polynomial_data, columns=columns)
            csv_filename = f'pce_coefficients_{specific_output}_degree_{degree}.csv' if specific_output else f'pce_coefficients_degree_{degree}.csv'
            coeff_df.to_csv(os.path.join(self.output_folder, csv_filename), index=False)
            print(f"\nCoefficients, orders, and polynomial terms saved to '{csv_filename}'")
        except Exception as e:
            print(f"Failed to save coefficients to CSV: {str(e)}")

        # Export to Excel
        try:
            excel_filename = f'pce_coefficients_{specific_output}_degree_{degree}.xlsx' if specific_output else f'pce_coefficients_degree_{degree}.xlsx'
            coeff_df.to_excel(os.path.join(self.output_folder, excel_filename), index=False, engine='openpyxl')
            print(f"Coefficients, orders, and polynomial terms saved to Excel: '{excel_filename}'")
        except Exception as e:
            print(f"Failed to save coefficients to Excel: {str(e)}")

        # Export to Word
        try:
            doc = Document()
            doc.add_heading(f"PCE Polynomial Coefficients - {specific_output or 'All Outputs'} (Degree {degree})", level=1)
            doc.add_paragraph(f"Generated by: Masud Babayev\nProject: Uncertainty Quantification in Engineering (PCE Project)\nDate: April 18, 2025\n")

            # Add table for coefficients and terms
            table = doc.add_table(rows=len(all_polynomial_data) + 1, cols=len(columns))
            table.style = 'Table Grid'

            # Add header row
            for col_idx, col_name in enumerate(columns):
                table.cell(0, col_idx).text = col_name

            # Add data rows
            for row_idx, data in enumerate(all_polynomial_data, 1):
                for col_idx, col_name in enumerate(columns):
                    table.cell(row_idx, col_idx).text = str(data[col_name])

            # Add the full polynomial as a paragraph
            doc.add_heading(f"Full Polynomial for {specific_output or 'All Outputs'}", level=2)
            full_poly = " + ".join(polynomial_terms) if polynomial_terms else "0"
            doc.add_paragraph(full_poly)

            # Save the Word document
            word_filename = f'pce_coefficients_{specific_output}_degree_{degree}.docx' if specific_output else f'pce_coefficients_degree_{degree}.docx'
            doc.save(os.path.join(self.output_folder, word_filename))
            print(f"Coefficients, orders, and polynomial terms saved to Word: '{word_filename}'")
        except Exception as e:
            print(f"Failed to save coefficients to Word: {str(e)}")

    def plot_input_distributions(self, df: pd.DataFrame):
        """Plot distributions of input parameters."""
        plt.figure(figsize=(15, 5))
        for i, param in enumerate(self.input_params, 1):
            plt.subplot(1, len(self.input_params), i)
            data = df[param].dropna()
            if data.empty:
                print(f"No data for {param} distribution plot")
                plt.text(0.5, 0.5, f"No Data for {param}", ha='center', va='center')
                continue
            sns.histplot(data, kde=True, color='blue')
            plt.title(f"{param} Distribution")
            plt.xlabel(param)
        plt.tight_layout()
        plt.savefig(os.path.join(self.output_folder, 'input_distributions.png'), dpi=300)
        plt.close()
        print("Input distributions plotted")

    def plot_pdf_cdf(self, mc_data, pce_data, output, degree):
        """Plot PDF and CDF comparisons for MC and PCE data."""
        valid_mc = mc_data[np.isfinite(mc_data)]
        valid_pce = pce_data[np.isfinite(pce_data)]
        if len(valid_mc) < 2 or len(valid_pce) < 2:
            print(f"Not enough data for {output} PDF/CDF at degree {degree}")
            return
        plt.figure(figsize=(10, 5))
        plt.subplot(1, 2, 1)
        sns.kdeplot(data=valid_mc, label='MC', color='blue')
        sns.kdeplot(data=valid_pce, label='PCE', color='orange')
        plt.title(f'{output} PDF (Degree {degree})')
        plt.legend()
        plt.subplot(1, 2, 2)
        if len(valid_mc) > 0 and len(valid_pce) > 0:
            mc_cdf = np.sort(valid_mc)
            pce_cdf = np.sort(valid_pce)
            plt.plot(mc_cdf, np.linspace(0, 1, len(mc_cdf)), label='MC', color='blue')
            plt.plot(pce_cdf, np.linspace(0, 1, len(pce_cdf)), label='PCE', color='orange')
            plt.title(f'{output} CDF (Degree {degree})')
            plt.legend()
        plt.tight_layout()
        plt.savefig(os.path.join(self.output_folder, f'pdf_cdf_{output}_degree_{degree}.png'), dpi=300)
        plt.close()

    def plot_residuals(self, mc_data, pce_data, output, degree):
        """Plot residuals (PCE - MC) for each output."""
        valid = np.isfinite(mc_data) & np.isfinite(pce_data)
        if valid.sum() < 2:
            print(f"Not enough data for {output} residuals at degree {degree}")
            return
        residuals = pce_data[valid] - mc_data[valid]
        plt.figure(figsize=(6, 6))
        plt.scatter(mc_data[valid], residuals, alpha=0.5, color='blue')
        plt.axhline(0, color='red', linestyle='--')
        plt.title(f'{output} Residuals (Degree {degree})')
        plt.xlabel('True Value')
        plt.ylabel('Residual')
        plt.grid(True)
        plt.savefig(os.path.join(self.output_folder, f'residuals_{output}_degree_{degree}.png'), dpi=300)
        plt.close()

    def plot_error_trend(self, errors):
        """Plot relative error (%) vs PCE degree for each output."""
        plt.figure(figsize=(10, 6))
        for output in self.outputs:
            degrees = sorted(errors[output].keys())
            if not degrees:
                continue
            rel_error = [errors[output][d] for d in degrees]
            plt.plot(degrees, rel_error, marker='o', label=output)
        plt.title('Relative Error (%) vs PCE Degree')
        plt.xlabel('PCE Degree')
        plt.ylabel('Relative Error (%)')
        plt.legend()
        plt.grid(True)
        plt.savefig(os.path.join(self.output_folder, 'error_trend.png'), dpi=300)
        plt.close()

    def plot_convergence(self, mc_means, pce_means):
        """Plot PCE mean convergence across degrees with MC mean."""
        plt.figure(figsize=(10, 6))
        for output in self.outputs:
            mc_mean = mc_means[output]
            degrees = sorted(pce_means[output].keys())
            if not degrees:
                continue
            pce_values = [pce_means[output][d] for d in degrees]
            plt.plot(degrees, pce_values, marker='o', label=output)
            plt.axhline(mc_mean, color='gray', linestyle='--', label=f'{output} MC Mean')
        plt.title('PCE Mean Convergence')
        plt.xlabel('PCE Degree')
        plt.ylabel('Mean Value')
        plt.legend()
        plt.grid(True)
        plt.savefig(os.path.join(self.output_folder, 'convergence.png'), dpi=300)
        plt.close()

    def plot_sensitivity(self, pce_models, joint_dist, degree):
        """Perform sensitivity analysis using Sobol indices and plot the results."""
        plt.figure(figsize=(10, 5))
        for i, (model, output) in enumerate(zip(pce_models, self.outputs), 1):
            sobol_indices = cp.Sens_m(model, joint_dist)
            total = sum(sobol_indices)
            if total > 0:
                sobol_indices = (sobol_indices / total) * 100
            else:
                sobol_indices = np.zeros_like(sobol_indices)

            plt.subplot(1, len(self.outputs), i)
            plt.bar(self.input_params, sobol_indices, color='teal')
            plt.title(f'Sensitivity for {output} (Degree {degree})')
            plt.xlabel('Input Parameter')
            plt.ylabel('Contribution to Variance (%)')
            plt.xticks(rotation=45)
            for j, v in enumerate(sobol_indices):
                plt.text(j, v + 1, f'{v:.1f}%', ha='center')

        plt.tight_layout()
        plt.savefig(os.path.join(self.output_folder, f'sensitivity_analysis_degree_{degree}.png'), dpi=300)
        plt.close()
        print(f"Sensitivity analysis plotted for degree {degree}")

    def run_analysis(self, df: pd.DataFrame) -> dict:
        """Run PCE analysis for orders 1 to 6, compare with MC, and extract coefficients at specified degrees."""
        X = df[self.input_params].values
        Y = np.array([df[out].values for out in self.outputs])

        self.plot_input_distributions(df)

        joint_dist = self.get_distributions()

        errors = {out: {} for out in self.outputs}
        pce_means = {out: {} for out in self.outputs}
        mc_means = {}
        pce_models_dict = {}
        poly_dict = {}

        for i, out in enumerate(self.outputs):
            valid_mc = Y[i][np.isfinite(Y[i])]
            mc_means[out] = np.mean(valid_mc) if len(valid_mc) >= 2 else np.nan

        best_order = None
        best_error = float('inf')
        best_watercut_order = None
        best_watercut_error = float('inf')

        for degree in range(1, self.max_pce_degree + 1):
            print(f"\nTesting PCE Degree {degree}...")
            try:
                pce_models, poly = self.build_pce_model(X, Y, joint_dist, degree)
                pce_pred = np.array([model(*X.T) for model in pce_models])
                pce_models_dict[degree] = pce_models
                poly_dict[degree] = poly

                avg_error = 0
                valid_outputs = 0
                for i, out in enumerate(self.outputs):
                    rel_error = self.calculate_relative_error(Y[i], pce_pred[i])
                    errors[out][degree] = rel_error
                    valid_pce = pce_pred[i][np.isfinite(pce_pred[i])]
                    pce_means[out][degree] = np.mean(valid_pce) if len(valid_pce) >= 2 else np.nan
                    if not np.isnan(rel_error):
                        avg_error += rel_error
                        valid_outputs += 1
                    if out == 'WaterCut' and not np.isnan(rel_error):
                        if rel_error <= best_watercut_error:
                            best_watercut_error = rel_error
                            best_watercut_order = degree

                    self.plot_pdf_cdf(Y[i], pce_pred[i], out, degree)
                    self.plot_residuals(Y[i], pce_pred[i], out, degree)

                if valid_outputs > 0:
                    avg_error /= valid_outputs
                    print(f"Average relative error (%) for degree {degree}: {avg_error:.2f}%")
                    if avg_error < best_error:
                        best_error = avg_error
                        best_order = degree

            except Exception as e:
                print(f"Degree {degree} failed: {str(e)}")
                print("Traceback:")
                traceback.print_exc()
                continue

        self.plot_error_trend(errors)
        self.plot_convergence(mc_means, pce_means)

        # Export errors to Excel
        print("\nExporting relative errors to Excel...")
        error_data = []
        for degree in range(1, self.max_pce_degree + 1):
            row = {'Degree': degree}
            for out in self.outputs:
                row[out] = errors[out].get(degree, np.nan)
            error_data.append(row)

        error_df = pd.DataFrame(error_data, columns=['Degree'] + self.outputs)
        error_excel_filename = os.path.join(self.output_folder, 'pce_relative_errors.xlsx')
        error_df.to_excel(error_excel_filename, index=False, engine='openpyxl')
        print(f"Relative errors saved to Excel: '{error_excel_filename}'")

        print("\nExtracting coefficients and polynomials at specified degrees...")
        if 6 in pce_models_dict:
            print("\nExtracting coefficients for WaterCut at degree 6...")
            watercut_model = [pce_models_dict[6][self.outputs.index('WaterCut')]]
            self.extract_coefficients_and_polynomial(watercut_model, poly_dict[6], 6, specific_output='WaterCut')

        if 4 in pce_models_dict:
            print("\nExtracting coefficients for OilRecoveryFactor at degree 4...")
            oilrecovery_model = [pce_models_dict[4][self.outputs.index('OilRecoveryFactor')]]
            self.extract_coefficients_and_polynomial(oilrecovery_model, poly_dict[4], 4, specific_output='OilRecoveryFactor')

        if 3 in pce_models_dict:
            print("\nExtracting coefficients for OIIP at degree 3...")
            oiip_model = [pce_models_dict[3][self.outputs.index('OIIP')]]
            self.extract_coefficients_and_polynomial(oiip_model, poly_dict[3], 3, specific_output='OIIP')

        if best_watercut_order is not None:
            print(f"\nUsing degree {best_watercut_order} for sensitivity analysis (best WaterCut error: {best_watercut_error:.2f}%)")
            self.plot_sensitivity(pce_models_dict[best_watercut_order], joint_dist, best_watercut_order)

        print("\n=== Analysis Results ===")
        for out in self.outputs:
            print(f"\n{out}:")
            for degree in range(1, self.max_pce_degree + 1):
                if degree in errors[out]:
                    rel_error = errors[out][degree]
                    print(f"Degree {degree} Relative Error: {rel_error:.2f}%")
                    if out == 'WaterCut' and not np.isnan(rel_error):
                        if rel_error < 15:
                            print(f"WaterCut relative error at degree {degree} is below 15%—great match!")
                            if degree > 1 and errors[out].get(degree - 1, float('inf')) <= rel_error + 0.01:
                                print(f"WaterCut error stabilized at degree {degree - 1} (no significant change)")
                        else:
                            print(f"WaterCut relative error at degree {degree} is above 15%—consider higher order.")

        print(f"\nBest PCE order (lowest average error): {best_order} with average relative error: {best_error:.2f}%")
        print(f"Best WaterCut relative error: {best_watercut_error:.2f}% at order {best_watercut_order}")
        print("Check the output_plots folder for visualizations.")

        return {
            'errors': errors,
            'best_order': best_order,
            'best_watercut_order': best_watercut_order,
            'success': best_order is not None
        }

def main():
    file_path = '/Users/masudbabayev/Library/CloudStorage/OneDrive-KFUPM/PhD/PhD Courses/Uncertainty Quantification in eng/PCE Project'
    input_file = '5000SAMPLE WITH PERMK.csv'

    try:
        pce_mc = PCEvsMC(file_path, input_file)
        df = pce_mc.load_data()
        results = pce_mc.run_analysis(df)

        if not results.get('success', False):
            print("\nAnalysis completed with errors.")
            return

    except Exception as e:
        print(f"\nFatal error in analysis: {str(e)}")
        raise

if __name__ == "__main__":
    main()